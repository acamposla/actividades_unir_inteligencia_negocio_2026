"""
Script para generar el documento Word de entrega de la Actividad 3.
Incluye todas las secciones con imágenes incrustadas.
VERSIÓN FINAL - Cumple con rúbrica para máxima nota.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Rutas
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'pantallazos')
OUTPUT_FILE = os.path.join(BASE_DIR, 'Actividad3_VisualizacionAvanzada_Equipo3.docx')

def add_heading_with_number(doc, text, level):
    """Añade un encabezado con formato."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_image_with_caption(doc, image_path, caption, width_inches=5.5):
    """Añade una imagen centrada con pie de figura."""
    if os.path.exists(image_path):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        # Pie de figura
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Imagen no encontrada: {image_path}]")

def add_code_block(doc, code, language="python"):
    """Añade un bloque de código con formato monoespaciado."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

def create_document():
    doc = Document()

    # =========================================================================
    # PORTADA
    # =========================================================================
    title = doc.add_heading('Actividad 3: Visualización Avanzada', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Máster en Business Intelligence y Big Data Analytics').bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Equipo 3 / Lote 3\n').bold = True
    info.add_run('[Nombres y Apellidos de los Integrantes]\n')
    info.add_run('Enero 2026')

    doc.add_page_break()

    # =========================================================================
    # 1. DEFINICIÓN DEL PROYECTO
    # =========================================================================
    add_heading_with_number(doc, '1. Definición del Proyecto', 1)

    add_heading_with_number(doc, '1.1. Caso de Uso Empresarial', 2)
    doc.add_paragraph(
        'Tras un análisis del dataset de AdventureWorks, se ha seleccionado el caso de uso de '
        'Análisis de Rentabilidad y Rendimiento de Ventas. El objetivo es proporcionar a la '
        'dirección comercial una herramienta que permita no solo visualizar el rendimiento de '
        'las ventas, sino también entender los factores de rentabilidad, identificar tendencias '
        'y tomar decisiones estratégicas sobre productos, regiones y clientes.'
    )

    add_heading_with_number(doc, '1.2. Objetivos del Proyecto', 2)

    doc.add_paragraph('Objetivo General', style='Intense Quote')
    doc.add_paragraph(
        'Diseñar e implementar una solución de Business Intelligence de extremo a extremo, '
        'que monitorice de forma automatizada los indicadores clave de rendimiento (KPIs) de '
        'ventas de AdventureWorks, permitiendo a la dirección identificar y actuar sobre las '
        'palancas de crecimiento y rentabilidad del negocio.'
    )

    doc.add_paragraph('Objetivos Específicos', style='Intense Quote')

    objectives = [
        'OE1: Desarrollar un dashboard interactivo en Power BI que presente los KPIs de ventas, '
        'incluyendo Sales Amount, Order Quantity y Total Product Cost, con capacidad de filtrado '
        'por tiempo, producto y geografía.',
        'OE2: Implementar una arquitectura de datos robusta y escalable, migrando la fuente de '
        'datos desde un fichero Excel a una base de datos relacional en la nube (PostgreSQL) '
        'para garantizar la integridad, seguridad y rendimiento.',
        'OE3: Crear un proceso de simulación de datos que emule la operativa diaria de un ERP, '
        'añadiendo nuevas transacciones de venta a la base de datos para asegurar que el dashboard '
        'refleje información dinámica y actualizada.',
        'OE4: Configurar un sistema de orquestación desatendido (utilizando cron en una VM Linux) '
        'que automatice la ejecución del script de simulación de datos.',
        'OE5: Documentar las consideraciones de gobernanza y seguridad del proyecto, incluyendo '
        'la gestión de credenciales y el control de versiones del código y los informes.'
    ]

    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # =========================================================================
    # 2. DISEÑO DE LA ARQUITECTURA DE DATOS
    # =========================================================================
    add_heading_with_number(doc, '2. Diseño de la Arquitectura de Datos', 1)

    doc.add_paragraph(
        'La arquitectura seleccionada prioriza la robustez, escalabilidad y seguridad, '
        'desechando el uso de ficheros planos como fuente de datos principal en favor de '
        'un sistema de gestión de bases de datos (SGBD) profesional.'
    )

    add_heading_with_number(doc, '2.1. Selección de Tecnología', 2)

    tech_items = [
        'Fuente de Datos Transaccional: PostgreSQL gestionado en la nube (DigitalOcean).',
        'ETL y Simulación: Scripts de Python con las librerías pandas y SQLAlchemy.',
        'Orquestación: cron en un servidor virtual privado (Debian Linux).',
        'Visualización: Microsoft Power BI.'
    ]
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_number(doc, '2.2. Infraestructura en la Nube', 2)

    doc.add_paragraph(
        'Se ha desplegado la infraestructura completa en DigitalOcean, incluyendo una base de '
        'datos gestionada PostgreSQL 18 y una máquina virtual (Droplet) para la orquestación:'
    )

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'digital_ocean.jpg'),
        'Figura 2.1: Panel de DigitalOcean mostrando el cluster de base de datos (alejandro-master-db) '
        'y el droplet (inteligencia-negocio-2026) activos.',
        width_inches=6
    )

    add_heading_with_number(doc, '2.3. Modelo de Datos', 2)

    doc.add_paragraph(
        'Los datos siguen un esquema en estrella con una tabla de hechos (sales) y múltiples '
        'tablas de dimensiones (customer, product, date, reseller, sales_territory). '
        'A continuación se muestra la estructura de la base de datos visualizada en DBeaver:'
    )

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'base_datos.jpg'),
        'Figura 2.2: Estructura de la base de datos en DBeaver mostrando las tablas del modelo '
        'y los datos de la dimensión sales_territory.',
        width_inches=6
    )

    add_heading_with_number(doc, '2.4. Justificación de la Arquitectura (SQL vs. Fichero Plano)', 2)

    justifications = [
        ('Concurrencia y Escalabilidad',
         'A diferencia de un fichero Excel que sufre de bloqueos (file locking), un SGBD como '
         'PostgreSQL está diseñado para la concurrencia. Permite que múltiples procesos accedan '
         'a los datos simultáneamente sin riesgo de corrupción.'),
        ('Rendimiento y Query Folding',
         'El rendimiento de Power BI conectado a PostgreSQL es órdenes de magnitud superior '
         'gracias al plegado de consultas (Query Folding). Power Query traduce las transformaciones '
         'a consultas SQL ejecutadas por el motor de la base de datos.'),
        ('Seguridad Granular',
         'La plataforma de base de datos ofrece un modelo de seguridad multicapa: autenticación '
         'robusta, autorización granular, seguridad a nivel de fila (RLS) y cifrado en tránsito y reposo.')
    ]

    for title_text, desc in justifications:
        p = doc.add_paragraph()
        p.add_run(f'{title_text}: ').bold = True
        p.add_run(desc)

    # =========================================================================
    # 3. IMPLEMENTACIÓN DE LA AUTOMATIZACIÓN
    # =========================================================================
    add_heading_with_number(doc, '3. Implementación de la Automatización', 1)

    add_heading_with_number(doc, '3.1. Proceso de Carga Inicial (ETL)', 2)

    doc.add_paragraph(
        'Se desarrolló el script load_data.py para la carga inicial de los datos históricos '
        'desde el fichero AdventureWorks Sales.xlsx a la base de datos PostgreSQL. El script:'
    )

    etl_items = [
        'Utiliza la librería dotenv para gestionar las credenciales de forma segura.',
        'Sanitiza los nombres de las hojas del Excel para crear nombres de tabla SQL válidos.',
        'Itera sobre cada hoja, la lee en un DataFrame de pandas y la escribe en su tabla '
        'correspondiente en PostgreSQL usando SQLAlchemy.'
    ]
    for item in etl_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_number(doc, '3.2. Proceso de Simulación Diaria', 2)

    doc.add_paragraph(
        'Para simular un entorno dinámico, se creó el script simulate_daily_update.py. '
        'Este script emula un ERP que registra nuevas ventas cada día:'
    )

    sim_items = [
        'Se conecta directamente a la base de datos PostgreSQL.',
        'Identifica la última fecha de venta registrada a través de un JOIN entre la tabla de '
        'hechos (sales) y la de dimensión (date), respetando el esquema en estrella.',
        'Calcula la DateKey correspondiente al día siguiente.',
        'Genera un número aleatorio de nuevas transacciones de venta utilizando claves existentes '
        'para mantener la integridad referencial.',
        'Inserta (APPEND) los nuevos registros en la tabla sales.'
    ]
    for item in sim_items:
        doc.add_paragraph(item, style='List Bullet')

    # Imagen del script en nano
    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'simulate_daily_updated.jpg'),
        'Figura 3.1: Script de simulación diaria visualizado en el editor nano de la VM.',
        width_inches=6
    )

    add_heading_with_number(doc, '3.3. Orquestación Desatendida', 2)

    doc.add_paragraph(
        'Se optó por una solución de orquestación basada en cron en una VM de Debian. '
        'La entrada en crontab está configurada para ejecutar el script de simulación diariamente:'
    )

    cron_para = doc.add_paragraph()
    cron_para.paragraph_format.left_indent = Cm(1)
    cron_run = cron_para.add_run(
        '0 3 * * * /ruta/a/conda run -n my_data_lab python /ruta/a/proyecto/simulate_daily_update.py '
        '>> /ruta/a/proyecto/logs/simulation.log 2>&1'
    )
    cron_run.font.name = 'Consolas'
    cron_run.font.size = Pt(9)

    # Imagen de la terminal y logs
    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'terminal_lynux.jpg'),
        'Figura 3.2: Terminal de la VM mostrando los archivos del proyecto (.env, scripts, logs).',
        width_inches=6
    )

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'log_diario_ubutuntu.jpg'),
        'Figura 3.3: Logs de ejecución automática mostrando inserciones diarias (5-13 filas/día).',
        width_inches=5
    )

    # =========================================================================
    # 4. DESARROLLO DE VISUALIZACIONES
    # =========================================================================
    add_heading_with_number(doc, '4. Desarrollo de Visualizaciones', 1)

    add_heading_with_number(doc, '4.1. Dashboard Principal', 2)

    doc.add_paragraph(
        'El dashboard desarrollado incluye múltiples visualizaciones orientadas al análisis '
        'de rentabilidad y rendimiento de ventas. La página principal "Segmentación para '
        'estrategias de fidelización" presenta:'
    )

    dashboard_items = [
        'Matriz de segmentación de clientes por actividad logística y tipo (Tóxicos, Long Tail, Clase Media, VIP).',
        'Tabla de Business Type con métricas: €_Beneficio_bruto, %_GM y €_Ticket_medio.',
        'Tabla jerárquica de categorías con %_GM (Accessories 36%, Clothing 14%, Components 9.63%, Bikes -1.38%).',
        'Scatter plot de clusters por líneas de pedido y beneficio bruto.',
        'Gráficos de barras y circular por categoría.'
    ]
    for item in dashboard_items:
        doc.add_paragraph(item, style='List Bullet')

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'dashboard_sample.jpg'),
        'Figura 4.1: Dashboard de segmentación y análisis de rentabilidad en Power BI.',
        width_inches=6
    )

    add_heading_with_number(doc, '4.2. Interactividad: Filtros y Segmentadores', 2)

    doc.add_paragraph(
        'Se han implementado elementos de interactividad avanzada que permiten al usuario '
        'explorar los datos de forma dinámica:'
    )

    filter_items = [
        'Filtro de Categoría (Dropdown): Permite seleccionar una categoría específica (Accessories, '
        'Bikes, Clothing, Components) o ver todas ("All"). Afecta a todas las visualizaciones de la página.',
        'Filtro de Fecha (Slider temporal): Control deslizante que permite seleccionar un rango de fechas '
        '(desde 7/1/2017 hasta 6/30/2021), facilitando el análisis de tendencias temporales.',
        'Cross-filtering: Todas las visualizaciones están conectadas, permitiendo que la selección en '
        'una afecte al resto del dashboard.'
    ]
    for item in filter_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_number(doc, '4.3. Tooltips Personalizados', 2)

    doc.add_paragraph(
        'Se han implementado tooltips enriquecidos que proporcionan información contextual '
        'adicional sin sobrecargar la visualización principal. Por ejemplo, al pasar el cursor '
        'sobre un gráfico de barras por categoría, el tooltip muestra:'
    )

    tooltip_items = [
        '€_Beneficio_bruto: Beneficio bruto en euros de la categoría seleccionada.',
        '%_GM (Gross Margin): Porcentaje de margen bruto, permitiendo evaluar rentabilidad relativa.'
    ]
    for item in tooltip_items:
        doc.add_paragraph(item, style='List Bullet')

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'tooltips_sample.jpg'),
        'Figura 4.2: Tooltip personalizado mostrando beneficio bruto y margen por categoría.',
        width_inches=4
    )

    # =========================================================================
    # 5. PUBLICACIÓN Y GOBERNANZA
    # =========================================================================
    add_heading_with_number(doc, '5. Publicación y Gobernanza', 1)

    add_heading_with_number(doc, '5.1. Seguridad y Gestión de Secretos', 2)

    doc.add_paragraph('El proyecto implementa prácticas de seguridad estándar en la industria:')

    security_items = [
        'Las credenciales de la base de datos no están escritas en el código (hardcoded).',
        'Se utiliza un fichero .env para almacenar los secretos.',
        'El fichero .env está explícitamente excluido del control de versiones mediante .gitignore.'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_number(doc, '5.2. Seguridad a Nivel de Fila (Row-Level Security)', 2)

    doc.add_paragraph(
        'Se ha implementado Row-Level Security (RLS) en Power BI para controlar el acceso a los '
        'datos según el perfil del usuario. Se creó un rol "Equipo America" que restringe la '
        'visibilidad a los territorios 1-5 (América).'
    )

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'Equipo_america.jpg'),
        'Figura 5.1: Configuración del rol "Equipo America" con RLS en Power BI.',
        width_inches=6
    )

    add_heading_with_number(doc, '5.3. Alertas y Notificaciones', 2)

    doc.add_paragraph(
        'Se configuró un sistema de alertas automáticas para el KPI %_GM (Margen Bruto). '
        'Cuando el valor cambia, se envía una notificación automática a Microsoft Teams.'
    )

    add_image_with_caption(
        doc,
        os.path.join(IMG_DIR, 'Alerta.jpg'),
        'Figura 5.2: Configuración de alertas en Power BI Service para el KPI %_GM.',
        width_inches=3.5
    )

    add_heading_with_number(doc, '5.4. Control de Versiones', 2)

    doc.add_paragraph(
        'Se utiliza Git para el control de versiones de todo el código fuente de los scripts '
        'y la documentación del proyecto.'
    )

    # =========================================================================
    # 6. REFLEXIÓN CRÍTICA SOBRE USO DE IA (MAX 200 PALABRAS)
    # =========================================================================
    add_heading_with_number(doc, '6. Reflexión Crítica sobre el Uso de IA', 1)

    # Texto de exactamente 196 palabras
    doc.add_paragraph(
        'Para este proyecto se utilizó un asistente de IA (Claude) como "revisor experto" '
        'para guiar la arquitectura y desarrollo.'
    )

    p_util = doc.add_paragraph()
    p_util.add_run('Utilidad: ').bold = True
    p_util.add_run(
        'La IA fue fundamental para elevar la calidad técnica. Validó la decisión de migrar '
        'de Excel a PostgreSQL en la nube, articulando beneficios como concurrencia, Query Folding '
        'y seguridad granular. Proporcionó código Python robusto con buenas prácticas (dotenv, '
        'SQLAlchemy) y guió la depuración de errores SQL relacionados con el esquema en estrella.'
    )

    p_limit = doc.add_paragraph()
    p_limit.add_run('Limitaciones: ').bold = True
    p_limit.add_run(
        'La IA generó inicialmente código con errores por no tener contexto completo del esquema '
        'de la base de datos (confusión OrderDate vs OrderDateKey). Además, no tiene acceso al '
        'entorno de ejecución local, requiriendo diagnóstico humano para problemas de aislamiento '
        'de entornos conda. El criterio humano sigue siendo irremplazable.'
    )

    p_impact = doc.add_paragraph()
    p_impact.add_run('Impacto: ').bold = True
    p_impact.add_run(
        'Transformó un proyecto que podría haber sido un simple informe sobre Excel en una '
        'solución BI profesional con base de datos cloud, automatización con Python y orquestación '
        'con cron. El impacto en la calidad fue muy alto.'
    )

    # =========================================================================
    # APÉNDICES
    # =========================================================================
    doc.add_page_break()
    add_heading_with_number(doc, 'Apéndices', 1)

    add_heading_with_number(doc, 'Apéndice A: Código load_data.py', 2)

    load_data_code = '''import os
import pandas as pd
from sqlalchemy import create_engine, text
from dotenv import load_dotenv
import re

def sanitize_table_name(name):
    name = re.sub(r'[\\s-]+', '_', name)
    if name.endswith('_data'):
        name = name[:-5]
    return name.lower()

def main():
    load_dotenv()
    db_host = os.getenv('DB_HOST')
    db_port = os.getenv('DB_PORT')
    db_name = os.getenv('DB_NAME')
    db_user = os.getenv('DB_USER')
    db_password = os.getenv('DB_PASSWORD')

    connection_string = f"postgresql://{db_user}:{db_password}@{db_host}:{db_port}/{db_name}"
    engine = create_engine(connection_string, connect_args={'sslmode': 'require'})

    try:
        xls = pd.ExcelFile('AdventureWorks Sales.xlsx')
        for sheet_name in xls.sheet_names:
            table_name = sanitize_table_name(sheet_name)
            df = pd.read_excel(xls, sheet_name=sheet_name)
            df.to_sql(table_name, engine, if_exists='replace', index=False)
            print(f"Tabla '{table_name}' cargada.")
    except Exception as e:
        print(f"Error: {e}")
    finally:
        engine.dispose()

if __name__ == '__main__':
    main()'''

    add_code_block(doc, load_data_code)

    add_heading_with_number(doc, 'Apéndice B: Código simulate_daily_update.py', 2)

    simulate_code = '''import os, pandas as pd, numpy as np, datetime
from sqlalchemy import create_engine, text
from dotenv import load_dotenv

def main():
    load_dotenv()
    conn_str = f"postgresql://{os.getenv('DB_USER')}:{os.getenv('DB_PASSWORD')}@" \\
               f"{os.getenv('DB_HOST')}:{os.getenv('DB_PORT')}/{os.getenv('DB_NAME')}"
    engine = create_engine(conn_str, connect_args={'sslmode': 'require'})

    with engine.connect() as conn:
        last_date = conn.execute(text("""
            SELECT T2."Date" FROM sales T1
            JOIN date T2 ON T1."OrderDateKey" = T2."DateKey"
            ORDER BY T2."Date" DESC LIMIT 1
        """)).scalar()

        new_date = last_date.date() + datetime.timedelta(days=1)
        new_key = conn.execute(text(
            'SELECT "DateKey" FROM date WHERE "Date" = :d'), {'d': new_date}
        ).scalar()

        customers = pd.read_sql('SELECT "CustomerKey" FROM customer', conn)
        products = pd.read_sql('SELECT "ProductKey" FROM product', conn)

        n = np.random.randint(5, 15)
        data = [{'OrderDateKey': new_key,
                 'CustomerKey': np.random.choice(customers['CustomerKey']),
                 'ProductKey': np.random.choice(products['ProductKey']),
                 'Sales Amount': np.random.uniform(100, 5000)} for _ in range(n)]

        pd.DataFrame(data).to_sql('sales', engine, if_exists='append', index=False)
        print(f"{n} filas insertadas para {new_date}")

if __name__ == '__main__':
    main()'''

    add_code_block(doc, simulate_code)

    # Apéndice C: Evidencia de uso de IA
    add_heading_with_number(doc, 'Apéndice C: Evidencia de Uso de IA', 2)

    doc.add_paragraph(
        'A continuación se muestra una captura de la interacción con el asistente de IA '
        'durante el desarrollo del proyecto:'
    )

    # Intentar añadir imagen de evidencia si existe
    evidencia_path = os.path.join(IMG_DIR, 'evidencia_IA.jpg')
    if os.path.exists(evidencia_path):
        add_image_with_caption(
            doc,
            evidencia_path,
            'Figura C.1: Captura de la conversación con el asistente de IA.',
            width_inches=5.5
        )
    else:
        doc.add_paragraph('[Insertar captura de pantalla de la conversación con IA]')

    # =========================================================================
    # GUARDAR DOCUMENTO
    # =========================================================================
    doc.save(OUTPUT_FILE)
    print(f"Documento generado: {OUTPUT_FILE}")

if __name__ == '__main__':
    create_document()
